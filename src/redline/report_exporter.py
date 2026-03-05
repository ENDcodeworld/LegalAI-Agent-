"""
Redline 报告导出器

支持导出格式：
- PDF：带红线圈注的 PDF 文档
- Word：带修订模式的 Word 文档
- HTML：网页版报告
- JSON：结构化数据
"""

import os
import json
from typing import Dict, List, Optional
from datetime import datetime
from dataclasses import asdict

from .models import RedlineReport, RedlineSuggestion, ClauseRedline, RedlineLevel, RiskSeverity


class RedlineReportExporter:
    """
    Redline 报告导出器
    
    支持多种格式导出，满足不同使用场景
    """
    
    def __init__(self, output_dir: str = "./reports"):
        """
        初始化导出器
        
        Args:
            output_dir: 输出目录
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export(
        self,
        report: RedlineReport,
        format: str = "pdf",
        filename: Optional[str] = None
    ) -> str:
        """
        导出报告
        
        Args:
            report: Redline 报告对象
            format: 导出格式 (pdf, word, html, json)
            filename: 输出文件名（不含扩展名）
            
        Returns:
            输出文件路径
        """
        if filename is None:
            filename = f"redline_report_{report.contract_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        if format == "json":
            return self._export_json(report, filename)
        elif format == "html":
            return self._export_html(report, filename)
        elif format == "pdf":
            return self._export_pdf(report, filename)
        elif format == "word":
            return self._export_word(report, filename)
        else:
            raise ValueError(f"不支持的导出格式：{format}")
    
    def _export_json(self, report: RedlineReport, filename: str) -> str:
        """导出为 JSON 格式"""
        filepath = os.path.join(self.output_dir, f"{filename}.json")
        
        data = report.to_dict()
        data['export_info'] = {
            'exported_at': datetime.now().isoformat(),
            'export_format': 'json',
            'version': '1.0'
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return filepath
    
    def _export_html(self, report: RedlineReport, filename: str) -> str:
        """导出为 HTML 格式"""
        filepath = os.path.join(self.output_dir, f"{filename}.html")
        
        html = self._generate_html_content(report)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        
        return filepath
    
    def _generate_html_content(self, report: RedlineReport) -> str:
        """生成 HTML 内容"""
        stats = report.get_statistics()
        
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Redline 报告 - {report.contract_name}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .container {{
            background: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
        }}
        h3 {{
            color: #7f8c8d;
        }}
        .summary-box {{
            background: #ecf0f1;
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 15px;
            margin: 20px 0;
        }}
        .stat-card {{
            background: white;
            padding: 15px;
            border-radius: 8px;
            text-align: center;
            border-left: 4px solid #3498db;
        }}
        .stat-value {{
            font-size: 24px;
            font-weight: bold;
            color: #2c3e50;
        }}
        .stat-label {{
            font-size: 14px;
            color: #7f8c8d;
        }}
        .risk-critical {{ border-left-color: #e74c3c; }}
        .risk-high {{ border-left-color: #e67e22; }}
        .risk-medium {{ border-left-color: #f1c40f; }}
        .risk-low {{ border-left-color: #27ae60; }}
        .clause-card {{
            background: #fafafa;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 20px;
            margin: 20px 0;
        }}
        .clause-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
        }}
        .clause-title {{
            font-size: 18px;
            font-weight: bold;
            color: #2c3e50;
        }}
        .risk-badge {{
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: bold;
            color: white;
        }}
        .badge-critical {{ background: #e74c3c; }}
        .badge-high {{ background: #e67e22; }}
        .badge-medium {{ background: #f1c40f; color: #333; }}
        .badge-low {{ background: #27ae60; }}
        .suggestion-card {{
            background: white;
            border: 1px solid #ddd;
            border-radius: 6px;
            padding: 15px;
            margin: 10px 0;
        }}
        .suggestion-header {{
            display: flex;
            justify-content: space-between;
            margin-bottom: 10px;
        }}
        .suggestion-level {{
            font-size: 12px;
            padding: 2px 8px;
            background: #3498db;
            color: white;
            border-radius: 4px;
        }}
        .original-text {{
            background: #fff3cd;
            padding: 10px;
            border-radius: 4px;
            margin: 10px 0;
            border-left: 3px solid #ffc107;
        }}
        .suggested-text {{
            background: #d4edda;
            padding: 10px;
            border-radius: 4px;
            margin: 10px 0;
            border-left: 3px solid #28a745;
        }}
        .rationale {{
            background: #e7f3ff;
            padding: 10px;
            border-radius: 4px;
            margin: 10px 0;
            font-size: 14px;
        }}
        .legal-basis {{
            font-size: 13px;
            color: #666;
            margin-top: 10px;
        }}
        .market-benchmark {{
            font-size: 13px;
            color: #666;
            margin-top: 5px;
        }}
        .confidence-bar {{
            width: 100px;
            height: 8px;
            background: #e0e0e0;
            border-radius: 4px;
            display: inline-block;
            overflow: hidden;
        }}
        .confidence-fill {{
            height: 100%;
            background: linear-gradient(90deg, #e74c3c, #f1c40f, #27ae60);
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        th, td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }}
        th {{
            background: #f8f9fa;
            font-weight: 600;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #999;
            font-size: 12px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📋 Redline 审阅报告</h1>
        
        <div class="summary-box">
            <h2>📄 合同信息</h2>
            <p><strong>合同名称：</strong>{report.contract_name}</p>
            <p><strong>报告编号：</strong>{report.report_id}</p>
            <p><strong>生成时间：</strong>{report.generated_at.strftime('%Y年%m月%d日 %H:%M')}</p>
        </div>
        
        <h2>📊 统计概览</h2>
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-value">{stats['total_suggestions']}</div>
                <div class="stat-label">总建议数</div>
            </div>
            <div class="stat-card risk-critical">
                <div class="stat-value">{stats['critical_count']}</div>
                <div class="stat-label">严重风险</div>
            </div>
            <div class="stat-card risk-high">
                <div class="stat-value">{stats['high_count']}</div>
                <div class="stat-label">高风险</div>
            </div>
            <div class="stat-card risk-medium">
                <div class="stat-value">{stats['medium_count']}</div>
                <div class="stat-label">中风险</div>
            </div>
            <div class="stat-card risk-low">
                <div class="stat-value">{stats['low_count']}</div>
                <div class="stat-label">低风险</div>
            </div>
            <div class="stat-card" style="border-left-color: #9b59b6;">
                <div class="stat-value">{stats['l4_suggestions']}</div>
                <div class="stat-label">L4 级建议</div>
            </div>
            <div class="stat-card" style="border-left-color: #1abc9c;">
                <div class="stat-value">{stats['l4_adoption_rate']}</div>
                <div class="stat-label">L4 占比</div>
            </div>
        </div>
        
        <h2>📝 执行摘要</h2>
        <div class="summary-box">
            <p>{report.executive_summary}</p>
        </div>
        
        <h2>⚠️ 整体风险评估</h2>
        <div class="summary-box" style="border-left: 4px solid {'#e74c3c' if '严重' in report.overall_risk_assessment else '#e67e22' if '高' in report.overall_risk_assessment else '#27ae60'};">
            <p>{report.overall_risk_assessment}</p>
        </div>
        
        <h2>💡 关键建议 (Top 5)</h2>
        <ol>
"""
        
        for rec in report.key_recommendations:
            html += f"            <li>{rec}</li>\n"
        
        html += """        </ol>
        
        <h2>📑 条款级 Redline 详情</h2>
"""
        
        for cr in report.clause_redlines:
            badge_class = {
                RiskSeverity.CRITICAL: 'badge-critical',
                RiskSeverity.HIGH: 'badge-high',
                RiskSeverity.MEDIUM: 'badge-medium',
                RiskSeverity.LOW: 'badge-low'
            }.get(cr.risk_level, 'badge-low')
            
            html += f"""
        <div class="clause-card">
            <div class="clause-header">
                <span class="clause-title">{cr.clause_title}</span>
                <span class="risk-badge {badge_class}">{cr.risk_level.value}</span>
            </div>
            
            <div class="original-text">
                <strong>原文：</strong><br>
                {cr.clause_content[:500]}{'...' if len(cr.clause_content) > 500 else ''}
            </div>
            
            {f'<div class="suggested-text"><strong>修改后：</strong><br>{cr.overall_suggestion[:500]}{"..." if len(cr.overall_suggestion) > 500 else ""}</div>' if cr.overall_suggestion else ''}
"""
            
            for s in cr.suggestions:
                level_badge = {
                    RedlineLevel.L4_INTELLIGENT: 'L4-智能 Redline',
                    RedlineLevel.L3_SPECIFIC: 'L3-具体建议',
                    RedlineLevel.L2_DIRECTION: 'L2-方向建议',
                    RedlineLevel.L1_GENERIC: 'L1-通用建议'
                }.get(s.level, 'L1-通用建议')
                
                html += f"""
            <div class="suggestion-card">
                <div class="suggestion-header">
                    <span><strong>建议 {s.suggestion_id}</strong> - {s.redline_type.value}</span>
                    <span class="suggestion-level">{level_badge}</span>
                </div>
                
                <div class="original-text">
                    <strong>❌ 原文：</strong><br>
                    {s.original_text}
                </div>
                
                <div class="suggested-text">
                    <strong>✅ 建议修改为：</strong><br>
                    {s.suggested_text}
                </div>
                
                <div class="rationale">
                    <strong>📖 修改理由：</strong><br>
                    {s.rationale}
                </div>
                
                <div style="margin-top: 10px; font-size: 13px;">
                    <strong>置信度：</strong>
                    <div class="confidence-bar">
                        <div class="confidence-fill" style="width: {s.confidence * 100}%"></div>
                    </div>
                    {s.confidence * 100:.0f}%
                    
                    <strong style="margin-left: 20px;">预计采纳率：</strong>
                    <div class="confidence-bar">
                        <div class="confidence-fill" style="width: {s.adoption_probability * 100}%"></div>
                    </div>
                    {s.adoption_probability * 100:.0f}%
                </div>
"""
                
                if s.legal_basis:
                    html += """                <div class="legal-basis">
                    <strong>⚖️ 法律依据：</strong><br>
"""
                    for lb in s.legal_basis:
                        html += f"""                    • {lb.law_name}{lb.article}：{lb.content}<br>
"""
                    html += """                </div>
"""
                
                if s.market_benchmark:
                    html += f"""                <div class="market-benchmark">
                    <strong>📊 市场基准：</strong>{s.market_benchmark.standard_practice}（采用率：{s.market_benchmark.adoption_rate * 100:.0f}%）
                </div>
"""
                
                html += """            </div>
"""
            
            html += """        </div>
"""
        
        html += f"""
        <div class="footer">
            <p>本报告由 LegalAI 智能 Redline 系统自动生成 | 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <p>注意：本报告仅供参考，不构成法律意见。重大合同请咨询专业律师。</p>
        </div>
    </div>
</body>
</html>
"""
        
        return html
    
    def _export_pdf(self, report: RedlineReport, filename: str) -> str:
        """
        导出为 PDF 格式
        
        使用 reportlab 或 weasyprint 生成 PDF
        这里使用 HTML 转 PDF 的方式
        """
        try:
            # 尝试使用 weasyprint
            from weasyprint import HTML
            
            html_content = self._generate_html_content(report)
            html_doc = HTML(string=html_content, encoding='utf-8')
            
            filepath = os.path.join(self.output_dir, f"{filename}.pdf")
            html_doc.write_pdf(filepath)
            
            return filepath
            
        except ImportError:
            # 如果 weasyprint 不可用，回退到 HTML 格式
            print("警告：weasyprint 未安装，无法生成 PDF。使用 HTML 格式代替。")
            print("安装命令：pip install weasyprint")
            return self._export_html(report, filename)
    
    def _export_word(self, report: RedlineReport, filename: str) -> str:
        """
        导出为 Word 格式（带修订模式）
        
        使用 python-docx 生成 Word 文档
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_UNDERLINE
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(f'Redline 审阅报告 - {report.contract_name}', 0)
            
            # 基本信息
            doc.add_heading('合同信息', level=1)
            doc.add_paragraph(f'报告编号：{report.report_id}')
            doc.add_paragraph(f'生成时间：{report.generated_at.strftime("%Y年%m月%d日 %H:%M")}')
            
            # 统计概览
            doc.add_heading('统计概览', level=1)
            stats = report.get_statistics()
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            headers = ['指标', '数值']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            data_rows = [
                ('总建议数', str(stats['total_suggestions'])),
                ('严重风险', str(stats['critical_count'])),
                ('高风险', str(stats['high_count'])),
                ('中风险', str(stats['medium_count'])),
                ('低风险', str(stats['low_count'])),
                ('L4 级建议', str(stats['l4_suggestions'])),
                ('L4 占比', stats['l4_adoption_rate']),
            ]
            
            for metric, value in data_rows:
                row = table.add_row()
                row.cells[0].text = metric
                row.cells[1].text = value
            
            # 执行摘要
            doc.add_heading('执行摘要', level=1)
            doc.add_paragraph(report.executive_summary)
            
            # 整体风险评估
            doc.add_heading('整体风险评估', level=1)
            doc.add_paragraph(report.overall_risk_assessment)
            
            # 关键建议
            doc.add_heading('关键建议 (Top 5)', level=1)
            for i, rec in enumerate(report.key_recommendations, 1):
                doc.add_paragraph(f'{i}. {rec}', style='List Number')
            
            # 条款级详情
            doc.add_heading('条款级 Redline 详情', level=1)
            
            for cr in report.clause_redlines:
                doc.add_heading(cr.clause_title, level=2)
                
                # 原文
                p = doc.add_paragraph()
                runner = p.add_run('原文：')
                runner.bold = True
                p.add_run(cr.clause_content[:500])
                
                # 修改后
                if cr.overall_suggestion:
                    p = doc.add_paragraph()
                    runner = p.add_run('修改后：')
                    runner.bold = True
                    runner.font_color.rgb = RGBColor(0, 128, 0)
                    p.add_run(cr.overall_suggestion[:500])
                
                # 各条建议
                for s in cr.suggestions:
                    doc.add_paragraph(f'建议 {s.suggestion_id} ({s.redline_type.value})', style='Heading 3')
                    
                    # 原文（删除线）
                    p = doc.add_paragraph()
                    runner = p.add_run('❌ 原文：')
                    runner.bold = True
                    runner = p.add_run(s.original_text)
                    runner.font_strike = True
                    runner.font_color.rgb = RGBColor(255, 0, 0)
                    
                    # 修改后（下划线）
                    p = doc.add_paragraph()
                    runner = p.add_run('✅ 建议：')
                    runner.bold = True
                    runner.font_color.rgb = RGBColor(0, 128, 0)
                    runner = p.add_run(s.suggested_text)
                    runner.underline = WD_UNDERLINE.SINGLE
                    runner.font_color.rgb = RGBColor(0, 128, 0)
                    
                    # 理由
                    p = doc.add_paragraph()
                    runner = p.add_run('📖 理由：')
                    runner.bold = True
                    p.add_run(s.rationale)
                    
                    # 法律依据
                    if s.legal_basis:
                        p = doc.add_paragraph()
                        runner = p.add_run('⚖️ 法律依据：')
                        runner.bold = True
                        for lb in s.legal_basis:
                            doc.add_paragraph(f'  • {lb.law_name}{lb.article}：{lb.content}', style='List Bullet')
                    
                    # 市场基准
                    if s.market_benchmark:
                        p = doc.add_paragraph()
                        runner = p.add_run('📊 市场基准：')
                        runner.bold = True
                        p.add_run(f'{s.market_benchmark.standard_practice}（采用率：{s.market_benchmark.adoption_rate * 100:.0f}%）')
            
            # 保存
            filepath = os.path.join(self.output_dir, f"{filename}.docx")
            doc.save(filepath)
            
            return filepath
            
        except ImportError:
            # 如果 python-docx 不可用，回退到 HTML 格式
            print("警告：python-docx 未安装，无法生成 Word 文档。使用 HTML 格式代替。")
            print("安装命令：pip install python-docx")
            return self._export_html(report, filename)
    
    def export_all_formats(
        self,
        report: RedlineReport,
        filename: Optional[str] = None
    ) -> Dict[str, str]:
        """
        导出所有格式
        
        Returns:
            各格式文件路径的字典
        """
        results = {}
        
        for format in ['json', 'html', 'pdf', 'word']:
            try:
                filepath = self.export(report, format, filename)
                results[format] = filepath
                print(f"✓ 已导出 {format.upper()} 格式：{filepath}")
            except Exception as e:
                print(f"✗ {format.upper()} 格式导出失败：{e}")
                results[format] = None
        
        return results


# 测试代码
if __name__ == "__main__":
    # 导入测试数据
    import sys
    sys.path.insert(0, '/home/admin/.openclaw/workspace/LegalAI-Agent/src')
    
    from analyzer.risk_analyzer import RiskAnalyzer
    from parser.contract_parser import ContractParser
    from .redline_generator import SmartRedlineGenerator
    
    # 示例合同
    sample_contract = """
    买卖合同
    
    甲方：北京科技有限公司
    乙方：上海贸易有限公司
    
    第一条 价格调整
    甲方有权随时调整产品价格，无需通知乙方。
    
    第二条 违约责任
    乙方违约需承担无限连带责任，赔偿甲方一切损失。
    
    第三条 争议解决
    本合同最终解释权归甲方所有。
    """
    
    # 解析和分析
    parser = ContractParser()
    contract = parser.parse_text(sample_contract)
    
    analyzer = RiskAnalyzer()
    analysis_result = analyzer.analyze(contract)
    
    # 生成 Redline
    generator = SmartRedlineGenerator()
    report = generator.generate_report(
        contract=contract,
        risk_points=analysis_result.risk_points,
        clauses=contract.clauses
    )
    
    # 导出报告
    exporter = RedlineReportExporter(output_dir='./test_reports')
    
    print("\n导出报告...")
    results = exporter.export_all_formats(report, filename='test_contract')
    
    print("\n导出完成:")
    for format, filepath in results.items():
        if filepath:
            print(f"  {format.upper()}: {filepath}")
